#!/usr/bin/env python3
"""
Create a professional Final Year Internship Weekly Report template
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border_to_paragraph(paragraph):
    """Add a border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    
    pPr.append(pBdr)

def create_weekly_report():
    """Create a professional weekly report template"""
    
    doc = Document()
    
    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('FINAL YEAR INTERNSHIP WEEKLY REPORT')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Add spacing
    doc.add_paragraph()
    
    # Create header information table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)
    
    # Fill in header fields
    headers = [
        ('Year:', ''),
        ('Week No.:', ''),
        ('Student Name and Enrolment Number:', ''),
        ('Supervisor Name:', ''),
        ('Project Title:', '')
    ]
    
    for i, (label, value) in enumerate(headers):
        row = table.rows[i]
        
        # Label cell
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Value cell (empty for user to fill)
        value_cell = row.cells[1]
        value_para = value_cell.paragraphs[0]
        value_para.text = value
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Section 1: WORK DONE
    section1_heading = doc.add_paragraph()
    section1_run = section1_heading.add_run('1. WORK DONE')
    section1_run.bold = True
    section1_run.font.size = Pt(12)
    section1_run.font.color.rgb = RGBColor(0, 51, 102)
    
    instruction1 = doc.add_paragraph()
    instruction1_run = instruction1.add_run('[Please write the details of the work done in the last fortnight]')
    instruction1_run.italic = True
    instruction1_run.font.size = Pt(10)
    instruction1_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add lines for writing
    for _ in range(20):
        p = doc.add_paragraph()
        p.add_run('_' * 100)
        p.paragraph_format.space_after = Pt(6)
    
    # Add page break
    doc.add_page_break()
    
    # Section 2: WORK TO BE DONE
    section2_heading = doc.add_paragraph()
    section2_run = section2_heading.add_run('2. WORK TO BE DONE')
    section2_run.bold = True
    section2_run.font.size = Pt(12)
    section2_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Add lines for writing
    for _ in range(12):
        p = doc.add_paragraph()
        p.add_run('_' * 100)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Section 3: PROBLEMS ENCOUNTERED
    section3_heading = doc.add_paragraph()
    section3_run = section3_heading.add_run('3. PROBLEMS ENCOUNTERED')
    section3_run.bold = True
    section3_run.font.size = Pt(12)
    section3_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Add lines for writing
    for _ in range(12):
        p = doc.add_paragraph()
        p.add_run('_' * 100)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Section 4: SELF EVALUATION OF THE PROGRESS
    section4_heading = doc.add_paragraph()
    section4_run = section4_heading.add_run('4. SELF EVALUATION OF THE PROGRESS')
    section4_run.bold = True
    section4_run.font.size = Pt(12)
    section4_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Add lines for writing
    for _ in range(18):
        p = doc.add_paragraph()
        p.add_run('_' * 100)
        p.paragraph_format.space_after = Pt(6)
    
    # Add footer with signatures
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature section
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.columns[0].width = Inches(3.25)
    sig_table.columns[1].width = Inches(3.25)
    
    # Student signature
    student_cell = sig_table.rows[0].cells[0]
    student_para = student_cell.paragraphs[0]
    student_para.add_run('_' * 30)
    student_para.add_run('\n')
    student_run = student_para.add_run('Student Signature')
    student_run.bold = True
    
    # Supervisor signature
    supervisor_cell = sig_table.rows[0].cells[1]
    supervisor_para = supervisor_cell.paragraphs[0]
    supervisor_para.add_run('_' * 30)
    supervisor_para.add_run('\n')
    supervisor_run = supervisor_para.add_run('Supervisor Signature')
    supervisor_run.bold = True
    supervisor_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save document
    output_file = '/Users/satyajeetsinhrathod/Desktop/Weekly_Report_Template.docx'
    doc.save(output_file)
    print(f"✅ Successfully created: {output_file}")

if __name__ == "__main__":
    create_weekly_report()
