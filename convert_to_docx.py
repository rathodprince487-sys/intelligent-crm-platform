#!/usr/bin/env python3
"""
Convert PROJECT_PROPOSAL.md to a well-formatted Word document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_formatted_text(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add formatted text to a paragraph"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def process_markdown_line(doc, line, in_code_block, code_lines):
    """Process a single line of markdown and add to document"""
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block - add all code lines
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'Code'
                for code_line in code_lines:
                    p.add_run(code_line + '\n')
                code_lines.clear()
            return False, code_lines
        else:
            # Start of code block
            return True, code_lines
    
    if in_code_block:
        code_lines.append(line)
        return True, code_lines
    
    # Skip empty lines in certain contexts
    if not line.strip():
        doc.add_paragraph()
        return False, code_lines
    
    # Handle horizontal rules
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.add_run('_' * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return False, code_lines
    
    # Handle headers
    if line.startswith('#'):
        level = len(re.match(r'^#+', line).group())
        text = line.lstrip('#').strip()
        
        if level == 1:
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p = doc.add_heading(text, level=1)
        elif level == 3:
            p = doc.add_heading(text, level=2)
        else:
            p = doc.add_heading(text, level=3)
        
        return False, code_lines
    
    # Handle bullet points
    if re.match(r'^\s*[-*]\s+', line):
        text = re.sub(r'^\s*[-*]\s+', '', line)
        p = doc.add_paragraph(text, style='List Bullet')
        return False, code_lines
    
    # Handle numbered lists
    if re.match(r'^\s*\d+\.\s+', line):
        text = re.sub(r'^\s*\d+\.\s+', '', line)
        p = doc.add_paragraph(text, style='List Number')
        return False, code_lines
    
    # Handle tables (basic support)
    if '|' in line and line.strip().startswith('|'):
        # Skip table separator lines
        if re.match(r'^\s*\|[\s\-:]+\|', line):
            return False, code_lines
        # For actual table rows, just add as normal paragraph for now
        # (Full table support would require more complex logic)
        p = doc.add_paragraph(line.strip())
        return False, code_lines
    
    # Handle bold and italic text in paragraphs
    p = doc.add_paragraph()
    
    # Process inline formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            text = part[2:-2]
            add_formatted_text(p, text, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            text = part[1:-1]
            add_formatted_text(p, text, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            # Inline code
            text = part[1:-1]
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Normal text
            p.add_run(part)
    
    return False, code_lines

def create_docx_from_markdown(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set up heading styles
    heading1 = doc.styles['Heading 1']
    heading1.font.size = Pt(24)
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1.font.bold = True
    
    heading2 = doc.styles['Heading 2']
    heading2.font.size = Pt(18)
    heading2.font.color.rgb = RGBColor(0, 102, 204)
    heading2.font.bold = True
    
    heading3 = doc.styles['Heading 3']
    heading3.font.size = Pt(14)
    heading3.font.color.rgb = RGBColor(51, 102, 153)
    heading3.font.bold = True
    
    # Create code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    except:
        # Style might already exist
        pass
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Process markdown
    in_code_block = False
    code_lines = []
    
    for line in lines:
        line = line.rstrip('\n')
        in_code_block, code_lines = process_markdown_line(doc, line, in_code_block, code_lines)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully created: {docx_file}")

if __name__ == "__main__":
    md_file = "/Users/satyajeetsinhrathod/n8n-backend/PROJECT_PROPOSAL.md"
    docx_file = "/Users/satyajeetsinhrathod/Desktop/PROJECT_PROPOSAL.docx"
    
    create_docx_from_markdown(md_file, docx_file)
